"""Export Markdown documents to Word (.docx) format using python-docx.

Handles the markdown elements produced by the generation agent:
  H1  → Heading 1 (centered, for the applicant's name in resumes)
  H2  → Heading 2 (section headers)
  H3  → Heading 3 (role titles / project names)
  -   → List Bullet
  1.  → List Number
  **  → bold inline
  *   → italic inline
  --- → horizontal separator (thin bottom border on an empty paragraph)
  plain text → Normal paragraph
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt as DPt, RGBColor


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def markdown_to_docx(markdown_text: str, output_path: str | Path) -> None:
    """Convert a Markdown string to a .docx file at output_path."""
    doc = Document()
    _configure_document(doc)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # H1
        if re.match(r"^# (?!#)", line):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_heading_font(p, 18, bold=True)

        # H2
        elif re.match(r"^## (?!#)", line):
            p = doc.add_heading(line[3:].strip(), level=2)
            _set_heading_font(p, 13, bold=True)
            _add_bottom_border(p)

        # H3
        elif re.match(r"^### (?!#)", line):
            p = doc.add_heading(line[4:].strip(), level=3)
            _set_heading_font(p, 11, bold=True)

        # Horizontal rule
        elif re.match(r"^[-*_]{3,}$", line):
            p = doc.add_paragraph()
            _add_bottom_border(p)

        # Bullet list
        elif re.match(r"^[-*+] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line[2:])

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\. ", "", line))

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)

        i += 1

    doc.save(str(output_path))


def export_document_set(
    docs,  # DocumentSet
    output_dir: str | Path,
) -> dict[str, Path]:
    """Export all present documents in a DocumentSet to DOCX. Returns {key: path}."""
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    written: dict[str, Path] = {}

    mapping = {
        "cover_letter": "cover_letter.docx",
        "resume": "resume.docx",
        "interview_guide": "interview_guide.docx",
    }
    for key, filename in mapping.items():
        content = docs.get(key)
        if content:
            path = out / filename
            markdown_to_docx(content, path)
            written[key] = path

    return written


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _configure_document(doc: Document) -> None:
    """Set sane default margins and base font."""
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = DPt(11)


def _set_heading_font(paragraph, size: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = DPt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def _add_bottom_border(paragraph) -> None:
    """Add a thin bottom border to a paragraph (acts as a section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# Regex to split inline text into bold/italic/plain segments
_INLINE_PATTERN = re.compile(
    r"(\*\*\*(?P<bi>.+?)\*\*\*"      # ***bold+italic***
    r"|\*\*(?P<b>.+?)\*\*"            # **bold**
    r"|\*(?P<i>.+?)\*"                # *italic*
    r"|`(?P<code>.+?)`)"              # `code`
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Parse inline markdown and add appropriately formatted runs."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])

        if m.group("bi"):
            run = paragraph.add_run(m.group("bi"))
            run.bold = True
            run.italic = True
        elif m.group("b"):
            run = paragraph.add_run(m.group("b"))
            run.bold = True
        elif m.group("i"):
            run = paragraph.add_run(m.group("i"))
            run.italic = True
        elif m.group("code"):
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Courier New"
            run.font.size = DPt(10)

        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])
