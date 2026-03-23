"""Tests for Markdown → DOCX export."""

import pytest
from docx import Document

from resume_refinery.exporters import export_document_set, markdown_to_docx


def test_markdown_to_docx_creates_file(tmp_path):
    md = "# John Smith\n\nSome cover letter text.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    assert out.exists()


def test_docx_has_heading_for_h1(tmp_path):
    md = "# Jordan Lee\n\nSome text.\n"
    out = tmp_path / "resume.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Jordan Lee" in p.text for p in headings)


def test_docx_has_bullet_list(tmp_path):
    md = "## Experience\n\n- Built something\n- Did another thing\n"
    out = tmp_path / "resume.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    bullet_paras = [p for p in doc.paragraphs if "List" in p.style.name]
    assert len(bullet_paras) >= 2


def test_docx_inline_bold(tmp_path):
    md = "I am a **bold statement** in text.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    # Find a run that is bold
    bold_runs = [
        run
        for para in doc.paragraphs
        for run in para.runs
        if run.bold and run.text.strip()
    ]
    assert any("bold statement" in r.text for r in bold_runs)


def test_export_document_set_writes_docx_files(tmp_path, document_set):
    written = export_document_set(document_set, tmp_path)
    assert "cover_letter" in written
    assert "resume" in written
    assert "interview_guide" in written
    assert written["cover_letter"].suffix == ".docx"
    assert written["cover_letter"].exists()


def test_export_document_set_partial(tmp_path):
    from resume_refinery.models import DocumentSet
    docs = DocumentSet(cover_letter="# Cover\n\nText.", resume=None, interview_guide=None)
    written = export_document_set(docs, tmp_path)
    assert "cover_letter" in written
    assert "resume" not in written


def test_markdown_to_docx_handles_horizontal_rule(tmp_path):
    md = "## Section\n\n---\n\nSome text after.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)  # Should not raise
    assert out.exists()


def test_docx_inline_italic(tmp_path):
    md = "This is *italic text* in a sentence.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    italic_runs = [
        run for para in doc.paragraphs for run in para.runs
        if run.italic and run.text.strip()
    ]
    assert any("italic text" in r.text for r in italic_runs)


def test_docx_inline_bold_italic(tmp_path):
    md = "This is ***bold and italic*** text.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    bi_runs = [
        run for para in doc.paragraphs for run in para.runs
        if run.bold and run.italic and run.text.strip()
    ]
    assert any("bold and italic" in r.text for r in bi_runs)


def test_docx_inline_code(tmp_path):
    md = "Use `kubectl apply` to deploy.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    code_runs = [
        run for para in doc.paragraphs for run in para.runs
        if run.font.name == "Courier New" and run.text.strip()
    ]
    assert any("kubectl apply" in r.text for r in code_runs)


def test_docx_numbered_list(tmp_path):
    md = "1. First item\n2. Second item\n3. Third item\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    numbered = [p for p in doc.paragraphs if "Number" in p.style.name]
    assert len(numbered) == 3


def test_docx_h3_heading(tmp_path):
    md = "### Senior Engineer @ DataFlow Inc\n\n- Built things\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert any("Senior Engineer" in p.text for p in h3)


def test_docx_skips_empty_lines(tmp_path):
    md = "# Name\n\n\n\nSome text.\n"
    out = tmp_path / "test.docx"
    markdown_to_docx(md, out)
    doc = Document(str(out))
    # Empty lines should not produce extra paragraphs
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty) == 2  # "Name" heading + "Some text."


def test_export_document_set_creates_output_dir(tmp_path):
    nested = tmp_path / "deep" / "nested" / "dir"
    from resume_refinery.models import DocumentSet
    docs = DocumentSet(cover_letter="# Letter\n\nBody.", resume="# Resume", interview_guide=None)
    written = export_document_set(docs, nested)
    assert nested.exists()
    assert "cover_letter" in written
    assert "resume" in written
    assert "interview_guide" not in written
